import pandas as pd
from docx import Document
import os

# Load customer data from Excel file
input_file = "Core.xlsx"  # Ensure this file is in the same directory
output_folder = "Generated_Word_Documents"

# Create output folder if not exists
os.makedirs(output_folder, exist_ok=True)

# Read Excel data
df = pd.read_excel(input_file)

# Function to generate Word document
def generate_word_document(account_no, unit_name, limit, address):
    doc = Document()

    # Add Unit Name
    doc.add_paragraph(f"Name of the Unit: {unit_name}")
    
    # Add Address if available
    if pd.notna(address):
        doc.add_paragraph(f"Address: {address}")
    else:
        doc.add_paragraph("Address: ")

    # Add Limit and Account No
    doc.add_paragraph(f"Limit: {limit}")
    doc.add_paragraph(f"ACCOUNT NO: {account_no}\n")

    # Bank Address and Letter
    doc.add_paragraph("The Branch Manager\nState Bank of India\nMalegaon 423203\n")
    doc.add_paragraph("Dear Sir,\nSTATEMENT OF STOCK AND BOOK DEBTS AS ON MONTH END")
    doc.add_paragraph("With reference to our / my Demand Cash Credit account we/ I give below the Summary of stock & book debts.")

    # Table Structure
    table = doc.add_table(rows=8, cols=6)
    table.style = "Table Grid"

    # Headers
    headers = ["Items", "Value of stocks", "Margin %", "Advance value (3-4)", "Sub-limit, if any", "Drawing Power"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Add Stock Categories
    stock_categories = [
        "(i) Raw materials",
        "(ii) Stores",
        "(iii) Stock-in-process",
        "(iv) Finished goods",
        "(v) Spares",
        "(vi) Book-debts including bills discounted",
        "Total of (i) to (vi)"
    ]

    for row_index, category in enumerate(stock_categories, start=1):
        table.cell(row_index, 0).text = category

    # Save the document
    file_name = os.path.join(output_folder, f"{unit_name}.docx")
    doc.save(file_name)
    print(f"Generated: {file_name}")

# Loop through all customers and generate Word documents
for _, row in df.iterrows():
    generate_word_document(row["Account No"], row["Unit Name"], row["Limit"], row.get("Address", ""))

print("✅ All Word documents generated successfully!")
