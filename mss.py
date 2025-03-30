import pandas as pd
import numpy as np
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Load customer data from Excel file
input_file = "MSS.xlsx"  # Ensure this file exists in the same directory
output_folder = "Generated_Mandatory_Stock_Statements"

# Create an output folder if it doesn't exist
os.makedirs(output_folder, exist_ok=True)

df_customers = pd.read_excel(input_file)

# Function to generate stock statement
def generate_stock_statement(customer_name, limit, address, account_no):
    months = ["April 2024", "May 2024", "June 2024", "July 2024", "August 2024", "September 2024", "October 2024", "November 2024", "December 2024", "January 2025", "February 2025", "March 2025"]
    
    # Generate stock statement data
    data = []
    for month in months:
        finished_goods = round(np.random.uniform(limit * 0.5, limit * 0.8), 2)
        book_debts = round(np.random.uniform(limit * 0.3, limit * 0.5), 2)
        dp = round(np.random.uniform(limit * 1.15, limit * 1.2), 2)  # DP is 15-20% more than the limit
        data.append([month, finished_goods, book_debts, dp])
    
    df = pd.DataFrame(data, columns=["Month", "Finished Goods", "Book Debts", "DP"])
    
    # Save to Excel
    excel_file = os.path.join(output_folder, f"{customer_name}.xlsx")
    df.to_excel(excel_file, index=False)
    
    print(f"Generated Excel: {excel_file}")
    
    # Generate Word documents
    for index, row in df.iterrows():
        generate_word_document(customer_name, address, account_no, row["Month"], row["Finished Goods"], row["Book Debts"], row["DP"], limit)

def generate_word_document(customer_name, address, account_no, month, finished_goods, book_debts, dp, limit):
    doc = Document()
    
    title = doc.add_paragraph()
    title_run = title.add_run(f"Stock Statement for {customer_name} - {month}")
    title_run.bold = True
    title_run.font.color.rgb = None  # Ensure the color is black
    title_run.font.size = Pt(14)
    
    # Create a small table for customer details
    details_table = doc.add_table(rows=4, cols=2)
    details_table.style = 'Table Grid'
    details_table.cell(0, 0).text = "Name of the Unit:"
    details_table.cell(0, 1).text = customer_name
    details_table.cell(1, 0).text = "Address:"
    details_table.cell(1, 1).text = address if pd.notna(address) else ""
    details_table.cell(2, 0).text = "Limit:"
    details_table.cell(2, 1).text = str(limit)
    details_table.cell(3, 0).text = "Account No:"
    details_table.cell(3, 1).text = str(account_no)
    
    doc.add_paragraph("\nTo")
    doc.add_paragraph("The Branch Manager")
    doc.add_paragraph("State Bank of India")
    doc.add_paragraph("Malegaon 423203")
    doc.add_paragraph("\nDear Sir,\n")
    doc.add_paragraph(f"STATEMENT OF STOCK AND BOOK DEBTS AS ON {month} END")
    
    # Create properly formatted table
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Items'
    hdr_cells[1].text = 'Value of Stocks'
    hdr_cells[2].text = 'Margin (%)'
    hdr_cells[3].text = 'Advance Value'
    hdr_cells[4].text = 'Sub-limit (if any)'
    hdr_cells[5].text = 'Drawing Power'
    
    # Data rows
    data = [
        ("Raw Materials", "", "", "", "", ""),
        ("Stock in Process", "", "", "", "", ""),
        ("Finished Goods", finished_goods, "25%", round(finished_goods * 0.75, 2), "-", "-"),
        ("Stores and Spares", "", "", "", "", ""),
        ("Book Debts", book_debts, "40%", round(book_debts * 0.6, 2), "-", "-"),
        ("Total DP", dp, "-", "-", "-", dp)
    ]
    
    for row_idx, (item, value, margin, advance, sub_limit, drawing_power) in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = item
        row_cells[1].text = str(value)
        row_cells[2].text = str(margin)
        row_cells[3].text = str(advance)
        row_cells[4].text = str(sub_limit)
        row_cells[5].text = str(drawing_power)
    
    # Save the Word document
    word_file = os.path.join(output_folder, f"{customer_name}_{month}.docx")
    doc.save(word_file)
    print(f"Generated Word: {word_file}")

# Run for all customers
for index, row in df_customers.iterrows():
    customer_name = row["Unit Name"]
    limit = row["Limit"]
    address = row.get("Address", "")  # Fetch address if available
    account_no = row.get("Account No", "")  # Fetch account number if available
    
    generate_stock_statement(customer_name, limit, address, account_no)

print("Stock statements and Word documents generated successfully!")
